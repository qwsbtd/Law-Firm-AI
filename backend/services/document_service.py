from pathlib import Path


def extract_text(file_path: str, mime_type: str) -> tuple[str, int]:
    """Return (full_text, page_count). Supports PDF, DOCX, TXT."""
    path = file_path.lower()

    if mime_type == "application/pdf" or path.endswith(".pdf"):
        import fitz  # PyMuPDF — package installs as pymupdf, imports as fitz
        doc = fitz.open(file_path)
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n\n".join(pages), len(pages)

    if (
        mime_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or path.endswith(".docx")
    ):
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        parts = []
        # Main body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Tables (billing records, fee schedules, etc.)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts), 1

    # Plain text fallback
    return Path(file_path).read_text(encoding="utf-8", errors="replace"), 1


def chunk_text(
    text: str,
    doc_id: int,
    filename: str,
    uploader_id: int,
    matter_id: int | None = None,
    matter_number: str = "",
) -> list[dict]:
    """
    Split text into ~512-token chunks (approximated as 400 words) with 38-word overlap.
    Returns list of dicts with 'text' and 'metadata' keys ready for embedding.
    All metadata values are str to satisfy ChromaDB requirements.
    """
    words = text.split()
    chunk_words = 400
    overlap_words = 38

    chunks = []
    start = 0
    chunk_index = 0
    while start < len(words):
        end = min(start + chunk_words, len(words))
        chunk_text_str = " ".join(words[start:end])
        chunks.append(
            {
                "text": chunk_text_str,
                "metadata": {
                    "doc_id": str(doc_id),
                    "filename": filename,
                    "uploader_id": str(uploader_id),
                    "chunk_index": str(chunk_index),
                    "matter_id": str(matter_id) if matter_id is not None else "",
                    "matter_number": matter_number,
                },
            }
        )
        chunk_index += 1
        if end == len(words):
            break
        start = end - overlap_words

    return chunks
